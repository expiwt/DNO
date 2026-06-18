#!/usr/bin/env python3
"""Generate a Word (.docx) report for the DNS + FNO project."""

import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.abspath(__file__))

# helpers

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p

def add_rich_para(doc, parts, size=11, spacing_after=6):
    """parts = list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing_after)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_image(doc, path, width=5.5):
    abs_path = os.path.join(ROOT, path)
    if os.path.exists(abs_path):
        doc.add_picture(abs_path, width=Inches(width))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return

# data

# Load leaderboard
with open(os.path.join(ROOT, "runs_navier_stokes_FiLM", "leaderboard.json")) as f:
    film_leaderboard = json.load(f)
with open(os.path.join(ROOT, "runs_navier_stokes_LpL", "bs32_m24_w32_lr0.001_st50", "logs", "summary.json")) as f:
    lpl_best = json.load(f)
with open(os.path.join(ROOT, "runs_navier_stokes_LpL", "bs32_m24_w32_lr0.0005_st50", "logs", "summary.json")) as f:
    lpl_second = json.load(f)

best_film = film_leaderboard[0]

# document

doc = Document()

# Title
title = doc.add_heading('Сравнение методов FiLM и LpL в FNO\nдля задачи Навье–Стокса (обратная ступенька)', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
title.runs[0].font.size = Pt(18)

add_para(doc, "NSS Lab — Отчёт по вычислительному эксперименту", italic=True, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()  # spacer

add_heading(doc, "1. Постановка задачи", level=1)

add_para(doc, (
    "Рассматривается задача моделирования течения вязкой несжимаемой жидкости "
    "в конфигурации «обратная ступенька» (backward-facing step) с помощью "
    "аппроксимации оператора решения уравнений Навье–Стокса нейронной сетью. "
    "Цель — по деформированным координатам расчётной сетки (X, Y) и числу "
    "Рейнольдса Re восстановить поля скорости (U, V) и давления (P)."
))

add_para(doc, (
    "Формально: дан оператор G, такой что G : (X, Y, Re) → (U, V, P). "
    "Нейронная сеть (FNO) обучается аппроксимировать G на наборе "
    "DNS-симуляций при разных Re."
))

add_para(doc, (
    "Вход: деформированные координаты X, Y (2 канала, сетка 128×128) + параметр Re."
))
add_para(doc, (
    "Выход: поля U, V, P (3 канала, сетка 128×128)."
))
add_para(doc, (
    "Диапазон Re: ~802 … 1353."
))

add_heading(doc, "2. Генерация данных (DNS)", level=1)

add_para(doc, (
    "Данные получены методом прямого численного моделирования (DNS) "
    "уравнений Навье–Стокса для несжимаемой жидкости в конфигурации "
    "«обратная ступенька» при различных числах Рейнольдса. "
    "Генерация выполнена на Python с использованием библиотеки Firedrake (FEM)."
))

# 2.1 Уравнение и дискретизация
add_heading(doc, "2.1. Уравнение и FEM-дискретизация", level=2)

add_para(doc, (
    "Решаются двумерные нестационарные уравнения Навье–Стокса для несжимаемой жидкости:"
))
add_para(doc, (
    "∂u/∂t + (u·∇)u = −∇p + ν∇²u,   ∇·u = 0"
), italic=True)

add_para(doc, (
    "Дискретизация по пространству — метод конечных элементов (FEM) с парой Тейлора–Худа "
    "P2–P1: скорость аппроксимируется квадратичными элементами (CG-2), "
    "давление — линейными (CG-1). Это inf-sup устойчивая пара, исключающая "
    "ложные осцилляции давления."
))

add_para(doc, (
    "По времени — неявная схема Эйлера (Backward Euler, fully implicit): "
    "конвективный член берётся на новом слое, что даёт безусловную устойчивость "
    "(отсутствие CFL-ограничения). Каждый шаг — решение одной большой "
    "разреженной системы через PETSc (Sparse LU или GMRES)."
))

# 2.2 Параметры симуляции
add_heading(doc, "2.2. Параметры симуляции", level=2)

add_para(doc, (
    "Для каждого из 709 семплов выполняется полный цикл DNS. Параметры:"
))

add_table(doc,
    ["Параметр", "Значение", "Описание"],
    [
        ["DT", "0.01", "Шаг по времени (безразмерный)"],
        ["T_BURN", "20.0", "Время «прогрева» — выход на стационарный режим (2000 шагов)"],
        ["T_SAMPLING", "25.0", "Время накопления статистики для усреднения (2500 шагов)"],
        ["Общее время", "45.0", "Итого 4500 временны́х шагов на один семпл"],
        ["Re", "800–1500", "Случайное равномерное распределение для каждого семпла"],
    ],
    col_widths=[1.5, 1.2, 4.0]
)

add_para(doc, (
    "Выбор T_BURN = 20.0 — это ≈20 convective time units (τ = H/U_avg). "
    "Этого достаточно, чтобы за ступенькой сформировалась стационарная зона "
    "рециркуляции и переходные процессы затухли. "
    "T_SAMPLING = 25.0 — интервал, на котором накапливается среднее."
))

# 2.3 Граничные условия
add_heading(doc, "2.3. Граничные условия", level=2)

add_table(doc,
    ["Маркер", "Граница", "Граничное условие"],
    [
        ["1", "Вход (inlet)", "Параболический профиль: u = [4·y·(H−y)/H², 0] (Poiseuille)"],
        ["3", "Нижняя стенка и ступенька", "No-slip: u = (0, 0)"],
        ["4", "Потолок (верхняя стенка)", "No-slip: u = (0, 0)"],
        ["—", "Выход (outlet)", "Do-nothing (natural BC из слабой формы)"],
    ],
    col_widths=[1.0, 2.5, 4.0]
)

add_para(doc, (
    "Высота канала H определяется из сетки для каждого семпла "
    "(конкретная геометрия ступеньки). "
    "Входной профиль — полностью развитое течение Пуазёйля."
))

# 2.4 Алгоритм генерации
add_heading(doc, "2.4. Алгоритм генерации (один семпл)", level=2)

add_para(doc, "Полный процесс для каждого из 709 семплов:", bold=True)

algo_items = [
    ("1. Загрузка сетки", (
        "Из папки train_domains_w_steps берётся .msh файл с уникальной геометрией "
        "ступеньки. Файлы отсортированы по номеру: step_1.msh … step_709.msh. "
        "Одновременно из x_data.csv / y_data.csv загружаются деформированные координаты "
        "128×128 — образ равномерной сетки при диффеоморфизме, отображающем "
        "прямоугольник в геометрию со ступенькой."
    )),
    ("2. Решение Стокса (начальное приближение)", (
        "Сначала решается уравнение Стокса ν∇²u = ∇p (без конвекции). "
        "Это даёт гладкое начальное поле, от которого стартует нестационарная симуляция."
    )),
    ("3. Нестационарный цикл NS", (
        "Запускается цикл на 4500 шагов (45.0 ед. времени). "
        "Первые 2000 шагов (T_BURN = 20.0) — выход потока на стационарный режим: "
        "за ступенькой формируется отрывной пузырь, переходные процессы затухают. "
        "Данные на этом этапе не записываются."
    )),
    ("4. Усреднение по времени", (
        "На интервале T_SAMPLING = 25.0 (шаги 2001–4500) поля скорости и давления "
        "накапливаются: u_sum += u(t), p_sum += p(t). "
        "После завершения цикла — деление на число шагов (2500): "
        "u_mean = u_sum / 2500, p_mean = p_sum / 2500. "
        "Это даёт осреднённое по времени стационарное решение."
    )),
    ("5. Интерполяция в сетку маппера", (
        "Осреднённые поля u_mean(x, y), p_mean(x, y), заданные на FEM-сетке, "
        "интерполируются в точки регулярной сетки 128×128 (valid_pts из x_data.csv). "
        "Точки, попавшие в твёрдое тело (ступеньку), помечаются как NaN — "
        "это маска геометрии."
    )),
    ("6. Сохранение", (
        "Результат: одна строка в каждом CSV (x, y, u, v, p, re). "
        "Строка содержит 16384 числа (128×128) — flattened 2D-поле. "
        "Также сохраняется PNG-визуализация для контроля."
    )),
]

for title, desc in algo_items:
    p = doc.add_paragraph(style='List Bullet')
    run_title = p.add_run(f"{title}: ")
    run_title.bold = True
    run_title.font.size = Pt(10.5)
    run_desc = p.add_run(desc)
    run_desc.font.size = Pt(10.5)

# 2.5 Геометрия и деформированные координаты
add_heading(doc, "2.5. Геометрия и деформированные координаты", level=2)

add_para(doc, (
    "Ключевая особенность датасета — использование деформированных координат (X, Y). "
    "Каждая из 709 сеток имеет разное положение и форму ступеньки. "
    "Диффеоморфизм отображает равномерный квадрат [0,1]² в физическое пространство "
    "со ступенькой. В результате X и Y — не декартовы координаты, а образ "
    "равномерной сетки при этом отображении. Это позволяет FNO учиться "
    "восстанавливать поля на сетках произвольной геометрии, сохраняя регулярную "
    "структуру тензора."
))

add_para(doc, (
    "Точки внутри твёрдого тела (ступеньки) содержат NaN в X и Y. "
    "При загрузке в модель они заменяются нулями. "
    "Маска геометрии восстанавливается по условию: точки, где X ≈ 0 и Y ≈ 0, "
    "находятся внутри твёрдого тела."
))

# 2.6 Характеристики датасета
add_heading(doc, "2.6. Характеристики датасета", level=2)

add_table(doc,
    ["Характеристика", "Значение"],
    [
        ["Количество семплов", "709"],
        ["Пространственная сетка", "128 × 128 (физически неравномерная, топологически равномерная)"],
        ["Диапазон Re", "800–1500 (равномерное случайное распределение)"],
        ["Среднее Re", "~1156"],
        ["Стандартное отклонение Re", "~197"],
        ["Выходные поля", "U, V (скорость), P (давление) — 3 канала"],
        ["Состояние потока", "Ламинарный / переходный (Re < 1500)"],
        ["Разбиение", "90% train / 10% test (≈638 / 71)"],
        ["Среднее значение U (физ.)", "~0.984"],
        ["Среднее значение V (физ.)", "~0.021"],
        ["Среднее значение P (физ.)", "~−0.365"],
        ["Нормализация", "Z-score (StandardScaler) — для каждого поля независимо"],
    ],
    col_widths=[2.5, 4.5]
)

add_para(doc, (
    "При загрузке (функция load_fluid_data) данные считываются из CSV, "
    "NaN заменяются нулями, после чего производится Z-score нормализация. "
    "Для FiLM Re остаётся скаляром [N, 1] и подаётся отдельно; "
    "для LpL Re размножается на пространство и становится третьим каналом "
    "размерности [N, H, W]."
))

add_heading(doc, "3. Архитектура и методы", level=1)

# 3.1 FNO
add_heading(doc, "3.1. Базовая архитектура FNO", level=2)
add_para(doc, (
    "В основе обоих методов лежит двумерный Fourier Neural Operator (FNO) "
    "с 6 спектральными свёрточными слоями (SpectralConv2d_fast). "
    "Каждый блок включает:"
))

items = [
    "Спектральную свёртку в пространстве Фурье (rfft2 → умножение на веса → irfft2);",
    "Skip-connection через Conv2d 1×1 (линейный канал);",
    "Геометрические bias от равномерного grid и от деформированной сетки (mesh) "
    "через Conv2d 1×1 (spatial conditioning);",
    "Активацию GELU.",
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10.5)

add_para(doc, (
    "После спектральных блоков — три полносвязных слоя (width → 256 → 128 → 128 → 3)."
))

add_para(doc, (
    "Количество мод Фурье (modes) и ширина скрытого пространства (width) — "
    "варьируемые гиперпараметры."
))

# 3.2 FiLM
add_heading(doc, "3.2. Метод FiLM (Feature-wise Linear Modulation)", level=2)
add_para(doc, (
    "В методе FiLM число Рейнольдса не подаётся на вход как канал признаков. "
    "Вместо этого Re модулирует скрытые представления внутри сети. "
    "На каждом из 6 спектральных блоков установлен FiLM-слой:"
))

add_para(doc, (
    "FiLM(re) → (γ, β) = MLP(re) → γ·x + β,  "
    "где γ, β — аффинные параметры (scale и shift), зависящие от Re."
), italic=True)

add_para(doc, (
    "FiLM-слой состоит из MLP: Linear(1 → 64) → GELU → Linear(64 → 2·width), "
    "где 2·width делится пополам на γ и β. "
    "Аргументация: такой подход позволяет Re влиять на все каналы признаков "
    "нелинейно, не «зашумляя» входной сигнал дополнительной координатой."
))

add_rich_para(doc, [
    ("Вход сети: ", True, False), ("[N, H, W, 2] — (X, Y) + Re как отдельный скаляр", False, False)
])

add_heading(doc, "3.2.1. Расшифровка [N, H, W, C]", level=3)

add_para(doc, (
    "Обозначение [N, H, W, C] описывает тензорные размерности данных, "
    "циркулирующих в нейронной сети:"
))

add_table(doc,
    ["Размерность", "Название", "В нашем случае"],
    [
        ["N", "Batch size", "Количество семплов в одном батче (16 или 32)"],
        ["H", "Height (высота)", "128 — пространственное разрешение по вертикали"],
        ["W", "Width (ширина)", "128 — пространственное разрешение по горизонтали"],
        ["C", "Channels (каналы)", "2 для FiLM (X, Y); 3 для LpL (X, Y, Re)"],
    ],
    col_widths=[1.2, 2.5, 3.5]
)

add_para(doc, (
    "Таким образом, тензор на входе FiLM имеет форму [16, 128, 128, 2] при батче 16, "
    "а для LpL — [16, 128, 128, 3]. После первого линейного слоя fc0 размерность "
    "каналов меняется на width (32 или 48), и остальные слои работают уже "
    "в скрытом пространстве. Свёртки в FNO выполняются по измерениям H и W, "
    "независимо для каждого канала. После спектральных блоков тензор имеет форму "
    "[N, H, W, width], а финальный fc5 проецирует его в [N, H, W, 3] — предсказание "
    "полей U, V, P."
))

# 3.3 LpL
add_heading(doc, "3.3. Метод LpL (Learned Parametric Learner)", level=2)
add_para(doc, (
    "В методе LpL Re подаётся как дополнительный канал на входе: "
    "входной тензор имеет размерность [N, H, W, 3] (X, Y, Re). "
    "Значение Re повторяется по всему пространству (одинаково для всех точек сетки). "
    "FiLM-слои отсутствуют — Re просто входит в сеть как признак наравне с координатами."
))

add_rich_para(doc, [
    ("Вход сети: ", True, False), ("[N, H, W, 3] — (X, Y, Re)", False, False)
])
add_rich_para(doc, [
    ("Различие: ", True, False), ("fc0 — Linear(3, width) вместо Linear(2, width); "
     "отсутствуют FiLM-слои; нет отдельного re_scalar на forward.", False, True)
])

# 3.4 Сравнение
add_heading(doc, "3.4. Ключевые различия", level=2)

add_table(doc,
    ["Характеристика", "FiLM", "LpL"],
    [
        ["Подача Re", "Через FiLM-слои (γ, β)", "Как 3-й канал на входе"],
        ["Входные каналы", "2 (X, Y) + Re-скаляр", "3 (X, Y, Re)"],
        ["FiLM-слои", "6 (на каждый блок)", "Нет"],
        ["Параметров (лучшая конфиг.)", "~28.4M", "~7.14M"],
        ["Гибкость модуляции", "Каждый блок адаптируется под Re", "Только на входе"],
    ],
    col_widths=[2.0, 2.5, 2.5]
)

add_heading(doc, "4. Эксперименты", level=1)

add_heading(doc, "4.1. Loss-функция", level=2)
add_para(doc, (
    "Используется относительная Lp-норма (Relative L2):"
))
add_para(doc, (
    "ℒ_rel(u, v) = ‖u − v‖₂ / ‖v‖₂"
), italic=True)
add_para(doc, (
    "Общая loss = ℒ_rel(U_pred, U_true) + ℒ_rel(V_pred, V_true) + ℒ_rel(P_pred, P_true). "
    "Loss считается для каждого поля отдельно и суммируется."
))

add_heading(doc, "4.2. Оптимизатор и планировщик", level=2)

add_table(doc,
    ["Параметр", "Значение"],
    [
        ["Оптимизатор", "AdamW"],
        ["Learning rate", "1e-3, 5e-4"],
        ["Weight decay", "1e-4"],
        ["Планировщик", "StepLR (step=50, gamma=0.5)"],
        ["Эпох", "120"],
        ["Batch size", "16, 32"],
        ["Seed", "42"],
        ["Активация", "GELU"],
    ],
    col_widths=[2.5, 3.5]
)

add_heading(doc, "4.3. Сетка гиперпараметров (sweep)", level=2)

add_table(doc,
    ["Гиперпараметр", "Значения"],
    [
        ["modes", "24, 32"],
        ["width", "32, 48"],
        ["batch_size", "16, 32"],
        ["learning_rate", "1e-3, 5e-4"],
        ["Всего комбинаций (FiLM)", "16 (из них выполнено 14)"],
    ],
    col_widths=[2.5, 3.5]
)

add_heading(doc, "5. Результаты", level=1)

add_heading(doc, "5.1. FiLM — Leaderboard (топ-5)", level=2)

film_rows = []
for r in film_leaderboard[:5]:
    film_rows.append([
        r["run"],
        r["best_test_loss"],
        r["params_count"],
        f"{r['time_sec']:.0f} сек"
    ])

add_table(doc,
    ["Run", "Best Test Loss", "Параметры", "Время"],
    film_rows,
    col_widths=[2.5, 1.5, 1.5, 1.5]
)

add_para(doc, "")
add_rich_para(doc, [
    ("Лучший FiLM: ", True, False),
    ("bs16_m32_w48_lr0.001_st50  |  test_loss = 0.03658  |  28.4M params", False, False)
])

add_heading(doc, "5.2. LpL — Результаты", level=2)

add_table(doc,
    ["Run", "Best Test Loss", "Параметры", "Время"],
    [
        ["bs32_m24_w32_lr0.001_st50", lpl_best["best_test_loss"], lpl_best["params_count"], f"{lpl_best['time_sec']:.0f} сек"],
        ["bs32_m24_w32_lr0.0005_st50", lpl_second["best_test_loss"], lpl_second["params_count"], f"{lpl_second['time_sec']:.0f} сек"],
    ],
    col_widths=[2.5, 1.5, 1.5, 1.5]
)

add_para(doc, "")
add_rich_para(doc, [
    ("Лучший LpL: ", True, False),
    ("bs32_m24_w32_lr0.001_st50  |  test_loss = 0.04070  |  7.14M params", False, False)
])

add_heading(doc, "5.3. Сравнение лучших конфигураций", level=2)

add_table(doc,
    ["Метод", "Конфигурация", "Test Loss", "Параметры", "Время"],
    [
        ["FiLM", "bs16_m32_w48", "0.03658", "28.43M", "~768 сек"],
        ["LpL",  "bs32_m24_w32", "0.04070", "7.14M",  "~458 сек"],
    ],
    col_widths=[1.0, 1.5, 1.2, 1.2, 1.2]
)

add_para(doc, "")

add_rich_para(doc, [
    ("Вывод: ", True, False),
    ("FiLM даёт на ~10% меньшую ошибку (0.0366 vs 0.0407), "
     "но требует в ~4 раза больше параметров (28.4M vs 7.1M) "
     "и на ~67% больше времени обучения. "
     "LpL — более лёгкая и быстрая альтернатива с конкурентной точностью.", False, False)
])

add_heading(doc, "5.4. Визуализация", level=2)

add_para(doc, (
    "Ниже приведён график сходимости loss для лучшего FiLM-ранах "
    "(bs16_m32_w48_lr0.001) и примеры предсказаний."
))

add_para(doc, "График loss (train + test):", bold=True)
add_image(doc, "runs_navier_stokes_FiLM/bs16_m32_w48_lr0.001_st50/plots/loss.png", width=5.0)

add_para(doc, "График loss для LpL (bs32_m24_w32_lr0.001):", bold=True)
add_image(doc, "runs_navier_stokes_LpL/bs32_m24_w32_lr0.001_st50/plots/loss.png", width=5.0)

add_para(doc, "Сравнение предсказаний FiLM vs LpL (пример Re≈1064):", bold=True)
add_image(doc, "runs_navier_stokes_FiLM/bs16_m32_w48_lr0.001_st50/test_inference/test_re1064_idx0.png", width=5.5)
add_image(doc, "runs_navier_stokes_LpL/bs32_m24_w32_lr0.001_st50/test_inference/test_re1064_idx0.png", width=5.5)

add_para(doc, "Сравнение предсказаний FiLM vs LpL (пример Re≈1288):", bold=True)
add_image(doc, "runs_navier_stokes_FiLM/bs16_m32_w48_lr0.001_st50/test_inference/test_re1288_idx35.png", width=5.5)
add_image(doc, "runs_navier_stokes_LpL/bs32_m24_w32_lr0.001_st50/test_inference/test_re1288_idx35.png", width=5.5)

add_heading(doc, "6. Заключение", level=1)

add_para(doc, (
    "В работе выполнено сравнение двух способов учёта параметра (Re) "
    "в операторном обучении для уравнений Навье–Стокса: явной FiLM-модуляции "
    "скрытых представлений против подачи параметра как дополнительного "
    "канала на вход (LpL)."
))

add_para(doc, "Основные результаты:")

concl = [
    "FiLM превосходит LpL по точности (loss 0.037 vs 0.041) за счёт "
    "адаптации каждого спектрального блока под конкретное Re.",
    "LpL значительно компактнее (7.1M против 28.4M параметров), "
    "быстрее обучается и при этом даёт приемлемую точность.",
    "Обучение при lr=1e-3 стабильно превосходит lr=5e-4 для обоих методов.",
    "Большее количество Фурье-мод (modes=32) систематически улучшает качество "
    "для FiLM; для LpL modes=24 также даёт хороший результат.",
]
for c in concl:
    p = doc.add_paragraph(c, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10.5)

# Save
out_path = os.path.join(os.path.expanduser("~/Desktop"), "DNS_FNO_Report.docx")
doc.save(out_path)
print(f" Report saved: {out_path}")
